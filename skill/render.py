#!/usr/bin/env python3
"""
Generate a FAANG-style monochrome resume (.docx + .pdf) from a YAML file.

Usage:
  python generate_faang.py
  python generate_faang.py --tags genai nlp mlops --out resume_genai
  python generate_faang.py --tags finance forecasting --out resume_banking
"""

import argparse
import yaml
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x00, 0x00, 0x00)
DARK = RGBColor(0x22, 0x22, 0x22)
FONT = "Georgia"          # classic, FAANG-typical serif; swap to "Calibri" for sans
CONTENT_WIDTH_IN = 7.1    # 8.5 - 2*0.7 margins


def keep(item_tags, wanted):
    return True if not wanted else any(t in wanted for t in (item_tags or []))


def _job_key(job):
    """Sort key: (year, month) of the start date, for reverse-chronological order."""
    try:
        mm, yy = str(job.get("start", "01/1900")).split("/")
        return (int(yy), int(mm))
    except Exception:
        return (0, 0)


def _experience(data):
    """Experience entries, always newest-first (deterministic — the model's
    ordering can't scramble the dates)."""
    return sorted(data.get("experience", []), key=_job_key, reverse=True)


def _clean(s):
    """Drop the em-dash 'AI tell' → plain hyphen. En-dash (date ranges) is kept."""
    return (str(s) if s is not None else "").replace(" — ", " - ").replace("—", "-").replace("―", "-")


def _sanitize(data):
    """Deep-clean em-dashes out of all string values in the resume dict."""
    if isinstance(data, str):
        return _clean(data)
    if isinstance(data, list):
        return [_sanitize(x) for x in data]
    if isinstance(data, dict):
        return {k: _sanitize(v) for k, v in data.items()}
    return data


def add_hrule(paragraph, color="000000", sz="4"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True   # never orphan a heading at page end
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK
    r.font.name = FONT
    # letter spacing
    rPr = r._element.get_or_add_rPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:val"), "30")
    rPr.append(spc)
    add_hrule(p)
    return p


def tabbed_line(doc, left_runs, right_text, right_italic=True, size=10):
    """One paragraph: left content, right-aligned date via right tab stop."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT)
    for txt, bold, italic in left_runs:
        r = p.add_run(txt)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = FONT
        r.font.color.rgb = DARK
    r = p.add_run("\t" + right_text)
    r.italic = right_italic
    r.font.size = Pt(size - 0.5)
    r.font.name = FONT
    r.font.color.rgb = DARK
    return p


# =========================================================================
# PDF renderer (ReportLab / Platypus) — pure Python, no external binary.
# Mirrors the DOCX style: monochrome, single-column, serif, right-aligned
# dates, section rules, experience-first ordering.
# =========================================================================

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                TableStyle, HRFlowable, KeepTogether)
from reportlab.lib.styles import ParagraphStyle
from xml.sax.saxutils import escape

PDF_BLACK = HexColor(0x000000)
PDF_DARK = HexColor(0x222222)
# Built-in PDF serif (Times) — no system-font dependency, so the renderer
# stays portable for a frozen single-file distribution (Task 4). The DOCX
# keeps Georgia; both are the serif look the design calls for.
SERIF = "Times-Roman"
SERIF_BOLD = "Times-Bold"
SERIF_ITALIC = "Times-Italic"

CONTENT_WIDTH = CONTENT_WIDTH_IN * inch
DATE_COL_W = 1.6 * inch    # right column reserved for the date


def _esc(text):
    return escape(str(text)) if text is not None else ""


def _date_row(left_html, right_text, story_styles):
    """A single line: left content (rich HTML), right-aligned date. Rendered as
    a borderless 2-col table so the date hugs the right margin like the DOCX."""
    left = Paragraph(left_html, story_styles["row_left"])
    right = Paragraph(_esc(right_text), story_styles["row_right"])
    t = Table([[left, right]], colWidths=[CONTENT_WIDTH - DATE_COL_W, DATE_COL_W])
    t.setStyle(TableStyle([
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
    ]))
    return t


def _pdf_styles():
    dark = "#222222"
    return {
        "name": ParagraphStyle("name", fontName=SERIF_BOLD, fontSize=20,
                               textColor=PDF_BLACK, alignment=TA_CENTER,
                               leading=23, spaceAfter=2),
        # splitLongWords=0 stops ReportLab from breaking a URL mid-token; the
        # line wraps only at the "  |  " separators, keeping each link intact.
        "contact": ParagraphStyle("contact", fontName=SERIF, fontSize=8.5,
                                  textColor=PDF_DARK, alignment=TA_CENTER,
                                  leading=11, spaceAfter=2, splitLongWords=0),
        "section": ParagraphStyle("section", fontName=SERIF_BOLD, fontSize=11,
                                  textColor=PDF_BLACK, leading=13, spaceBefore=8,
                                  spaceAfter=1),
        "summary": ParagraphStyle("summary", fontName=SERIF, fontSize=9.5,
                                  textColor=PDF_DARK, leading=12, spaceAfter=2),
        "row_left": ParagraphStyle("row_left", fontName=SERIF, fontSize=10.5,
                                   textColor=PDF_DARK, leading=13, alignment=TA_LEFT),
        # Same fontSize + leading as row_left so the right-aligned date shares the
        # exact baseline of the role heading in the 2-col table (VALIGN BOTTOM).
        "row_right": ParagraphStyle("row_right", fontName=SERIF_ITALIC, fontSize=10.5,
                                    textColor=PDF_DARK, leading=13, alignment=TA_RIGHT),
        "location": ParagraphStyle("location", fontName=SERIF_ITALIC, fontSize=8.5,
                                   textColor=PDF_DARK, leading=10, spaceAfter=1),
        "bullet": ParagraphStyle("bullet", fontName=SERIF, fontSize=9.5,
                                 textColor=PDF_DARK, leading=12, leftIndent=12,
                                 bulletIndent=2, spaceAfter=2),
        "skill": ParagraphStyle("skill", fontName=SERIF, fontSize=9.5,
                                textColor=PDF_DARK, leading=13.5, spaceAfter=5),
        "edu_sub": ParagraphStyle("edu_sub", fontName=SERIF_ITALIC, fontSize=9,
                                  textColor=PDF_DARK, leading=11, spaceAfter=3),
    }


def _pdf_section(story, styles, title):
    story.append(Paragraph(_esc(title).upper(), styles["section"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=PDF_BLACK,
                            spaceBefore=1, spaceAfter=3))


def build_pdf(data, wanted, out_stem):
    """Render the resume YAML straight to a .pdf with ReportLab. Returns path."""
    pdf_path = f"{out_stem}.pdf"
    styles = _pdf_styles()
    doc = SimpleDocTemplate(
        pdf_path, pagesize=letter,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        title=str(data.get("name", "Resume")),
    )
    story = []

    # Header
    story.append(Paragraph(_esc(data["name"]), styles["name"]))
    c = data["contact"]
    bits = [c.get("email"), c.get("phone"), c.get("location"),
            c.get("linkedin"), c.get("github"), c.get("medium")]
    bits = [_esc(b) for b in bits if b]
    if bits:
        story.append(Paragraph("  |  ".join(bits), styles["contact"]))

    # Summary
    if data.get("summary"):
        _pdf_section(story, styles, "Summary")
        story.append(Paragraph(_esc(" ".join(data["summary"].split())),
                               styles["summary"]))

    # Experience
    jobs = []
    for job in _experience(data):
        kept = [b for b in job.get("bullets", []) if keep(b.get("tags"), wanted)]
        if kept:
            jobs.append((job, kept))
    if jobs:
        _pdf_section(story, styles, "Experience")
        for job, kept in jobs:
            dates = f'{job.get("start","")} – {job.get("end","")}'
            left = (f'<b>{_esc(job["role"])}  -  </b>'
                    f'{_esc(job["company"])}')
            story.append(_date_row(left, dates, styles))
            if job.get("location"):
                story.append(Paragraph(_esc(job["location"]), styles["location"]))
            for b in kept:
                story.append(Paragraph(_esc(" ".join(b["text"].split())),
                                       styles["bullet"], bulletText="•"))

    # Skills
    groups = [g for g in data.get("skills", [])
              if keep(g.get("tags"), wanted) and g.get("items")]
    if groups:
        _pdf_section(story, styles, "Technical Skills")
        for g in groups:
            html = (f'<b>{_esc(g["category"])}:  </b>'
                    f'{_esc(" · ".join(g["items"]))}')
            story.append(Paragraph(html, styles["skill"]))

    # Education
    if data.get("education"):
        _pdf_section(story, styles, "Education")
        for e in data["education"]:
            dates = f'{e.get("start","")} – {e.get("end","")}'
            story.append(_date_row(f'<b>{_esc(e["degree"])}</b>', dates, styles))
            sub = _esc(e["institution"])
            if e.get("location"):
                sub += f'  |  {_esc(e["location"])}'
            story.append(Paragraph(sub, styles["edu_sub"]))

    doc.build(story)
    return pdf_path


def build_docx(data, wanted, out_stem):
    """Render the resume YAML to a .docx (python-docx). Returns the path."""
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.6)
        s.bottom_margin = Inches(0.6)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = DARK

    # Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(0)
    nr = name_p.add_run(data["name"])
    nr.bold = True
    nr.font.size = Pt(20)
    nr.font.name = FONT
    nr.font.color.rgb = BLACK
    rPr = nr._element.get_or_add_rPr()
    spc = OxmlElement("w:spacing"); spc.set(qn("w:val"), "40"); rPr.append(spc)

    c = data["contact"]
    bits = [c.get("email"), c.get("phone"), c.get("location"),
            c.get("linkedin"), c.get("github"), c.get("medium")]
    bits = [b for b in bits if b]
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(2)
    cr = cp.add_run("  |  ".join(bits))
    cr.font.size = Pt(8.5)
    cr.font.name = FONT
    cr.font.color.rgb = DARK

    # Summary
    if data.get("summary"):
        section_heading(doc, "Summary")
        sp = doc.add_paragraph(data["summary"].strip())
        sp.paragraph_format.space_after = Pt(3)
        for r in sp.runs:
            r.font.size = Pt(9.5); r.font.name = FONT; r.font.color.rgb = DARK

    # Experience
    jobs = []
    for job in _experience(data):
        kept = [b for b in job.get("bullets", []) if keep(b.get("tags"), wanted)]
        if kept:
            jobs.append((job, kept))
    if jobs:
        section_heading(doc, "Experience")
        for job, kept in jobs:
            dates = f'{job.get("start","")} \u2013 {job.get("end","")}'
            tabbed_line(doc,
                        [(job["role"] + "  -  ", True, False),
                         (job["company"], False, False)],
                        dates, size=10.5)
            if job.get("location"):
                loc_p = doc.add_paragraph()
                loc_p.paragraph_format.space_after = Pt(1)
                lr = loc_p.add_run(job["location"])
                lr.italic = True
                lr.font.size = Pt(8.5)
                lr.font.name = FONT
                lr.font.color.rgb = DARK
            for b in kept:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)
                bp.paragraph_format.space_before = Pt(0)
                br = bp.add_run(" ".join(b["text"].split()))
                br.font.size = Pt(9.5)
                br.font.name = FONT
                br.font.color.rgb = DARK

    # Skills
    groups = [g for g in data.get("skills", [])
              if keep(g.get("tags"), wanted) and g.get("items")]
    if groups:
        section_heading(doc, "Technical Skills")
        for g in groups:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)      # breathing room (was 2)
            p.paragraph_format.line_spacing = 1.15
            r1 = p.add_run(g["category"] + ":  ")
            r1.bold = True; r1.font.size = Pt(9.5); r1.font.name = FONT; r1.font.color.rgb = BLACK
            r2 = p.add_run(" · ".join(g["items"]))
            r2.font.size = Pt(9.5); r2.font.name = FONT; r2.font.color.rgb = DARK

    # Education \u2014 chain the whole section so it never splits across a page.
    if data.get("education"):
        edu_paras = [section_heading(doc, "Education")]
        for e in data["education"]:
            dates = f'{e.get("start","")} \u2013 {e.get("end","")}'
            dp = tabbed_line(doc, [(e["degree"], True, False)], dates, size=9.5)
            sub = e["institution"] + (f'  |  {e["location"]}' if e.get("location") else "")
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(3)
            sr = sp.add_run(sub)
            sr.italic = True; sr.font.size = Pt(9); sr.font.name = FONT; sr.font.color.rgb = DARK
            dp.paragraph_format.keep_together = True
            sp.paragraph_format.keep_together = True
            edu_paras += [dp, sp]
        for p in edu_paras[:-1]:
            p.paragraph_format.keep_with_next = True

    docx_path = f"{out_stem}.docx"
    doc.save(docx_path)
    return docx_path


def build(data, wanted, out_stem):
    """Render BOTH outputs from the same resume-YAML model, in pure Python.

    The .docx and .pdf are produced by two independent renderers (python-docx
    and ReportLab) — no DOCX→PDF conversion and no external binary. Returns
    (docx_path, pdf_path)."""
    data = _sanitize(data)
    docx_path = build_docx(data, wanted, out_stem)
    pdf_path = build_pdf(data, wanted, out_stem)
    return docx_path, pdf_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--yaml", default="store/career_history.yaml")
    ap.add_argument("--tags", nargs="*", default=None)
    ap.add_argument("--out", default="resume_faang")
    args = ap.parse_args()
    with open(args.yaml) as f:
        data = yaml.safe_load(f)
    wanted = set(args.tags) if args.tags else None
    d, p = build(data, wanted, args.out)
    print(d); print(p)


if __name__ == "__main__":
    main()
